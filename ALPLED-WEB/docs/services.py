import io
import json
import os
import time
import traceback
import hashlib
from pathlib import Path
from types import SimpleNamespace
from zipfile import BadZipFile, ZipFile
import requests
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode, urljoin
from urllib.request import Request, urlopen
from uuid import uuid4

from django.conf import settings
from django.db import transaction
from django.db.models import F, Max, Window
from django.db.models.functions import RowNumber
from django.urls import reverse
from docx import Document as DocxDocument

from common.models import Code
from common.onlyoffice import decode_jwt, encode_jwt
from common.project_selection import get_request_user
from common.storage import build_s3_uri, delete_object, delete_object_at_uri, read_bytes_from_uri, save_bytes
from files.models import ProjectFile
from projects.models import ProjectNet, ProjectUserRole

from .models import ApprovalReviewJob, Document, DocumentApproval, DocumentDetail, GenerationJob


DOCUMENT_CODE_SEQUENCE = ("DOC_SRS", "DOC_ITF", "DOC_ARCH", "DOC_ERD", "DOC_DB", "DOC_TS")
FILE_INPUT_DOCUMENT_CODES = {"DOC_SRS"}
DERIVED_DOCUMENT_CODES = {"DOC_ERD", "DOC_DB", "DOC_TS"}
DOCUMENT_PREREQUISITES = {
    "DOC_SRS": (),
    "DOC_ITF": ("DOC_SRS",),
    "DOC_ARCH": ("DOC_SRS",),
    "DOC_ERD": ("DOC_SRS",),
    "DOC_DB": ("DOC_SRS", "DOC_ERD"),
    "DOC_TS": ("DOC_SRS", "DOC_ITF"),
}

APPROVAL_STATUS_SEQUENCE = ("APRV_REQ", "APRV_COM", "APRV_RJT")
PROJECT_ROLE_CODES = ("ROLE_MANAGER", "ROLE_MEMBER")
GENERATION_SESSION_KEY = "docs_initial_generation"
DOC_JOB_SNAPSHOT_SESSION_KEY = "doc_job_snapshots"
ALLOWED_GENERATION_FILE_CODES = ("FILE_RFP", "FILE_MEETING")
INTERFACE_REFERENCE_DOCUMENT_CODE = "DOC_ITF"
ARCHITECTURE_DOCUMENT_CODE = "DOC_ARCH"
INTERFACE_REFERENCE_ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg"}
INTERFACE_REFERENCE_MAX_FILE_SIZE = 3 * 1024 * 1024
PROGRESS_PENDING = "PRGRS_PENDING"
PROGRESS_PROCESSING = "PRGRS_PROCESSING"
PROGRESS_COMPLETED = "PRGRS_COMPLETED"
PROGRESS_FAILED = "PRGRS_FAILED"
RUNNING_PROGRESS_CODES = (PROGRESS_PENDING, PROGRESS_PROCESSING)
TERMINAL_PROGRESS_CODES = (PROGRESS_COMPLETED, PROGRESS_FAILED)
FASTAPI_GENERATE_TIMEOUT_SECONDS = 10
FASTAPI_APPROVAL_REVIEW_TIMEOUT_SECONDS = 10
FASTAPI_CLIENT_USER_AGENT = "ALPLED-WEB/1.0 (Django; requests)"
WORKING_DOCUMENT_VERSIONS = ("0", "0.0")
GENERATION_JOB_KIND_INITIAL = "initial"
GENERATION_JOB_KIND_AUTO_APPLY = "auto_apply"


def _truncate_log_value(value, *, limit=600):
    text = str(value)
    if len(text) <= limit:
        return text
    return f"{text[:limit]}...<truncated>"


def _debug_generation_log(step, **fields):
    payload = {"step": step}
    payload.update(fields)
    try:
        message = json.dumps(payload, ensure_ascii=False, default=str)
    except TypeError:
        message = _truncate_log_value(payload)
    print(f"[docs.generate] {message}", flush=True)


def _summarize_generation_payload(payload):
    if not isinstance(payload, dict):
        return _truncate_log_value(payload)
    summary = dict(payload)
    summary["file_list_count"] = len(summary.get("file_list", []) or [])
    summary["image_list_count"] = len(summary.get("image_list", []) or [])
    if "file_list" in summary:
        summary["file_list"] = list(summary.get("file_list", []) or [])
    if "image_list" in summary:
        summary["image_list"] = list(summary.get("image_list", []) or [])
    return summary


def _build_empty_generation_state(project):
    return {
        "project_sn": project.sn if project else None,
        "selected_file_ids": [],
        "draft_documents": {},
        "confirmed_documents": {},
        "itf_reference_files": [],
    }


def get_doc_job_snapshots(session):
    snapshots = session.get(DOC_JOB_SNAPSHOT_SESSION_KEY, {})
    return snapshots if isinstance(snapshots, dict) else {}


def _normalize_doc_job_snapshot(payload):
    snapshot = dict(payload or {})
    snapshot["job_id"] = str(snapshot.get("job_id") or "").strip()
    snapshot["job_kind"] = str(snapshot.get("job_kind") or "").strip()
    snapshot["docs_cd"] = str(snapshot.get("docs_cd") or "").strip()
    snapshot["status"] = "running" if str(snapshot.get("status") or "").strip() in {"", "started", "running"} else str(snapshot.get("status")).strip()
    snapshot["tracking_document_sn"] = snapshot.get("tracking_document_sn")
    snapshot["job_status_code"] = str(snapshot.get("job_status_code") or PROGRESS_PENDING).strip()
    snapshot["job_status_label"] = str(snapshot.get("job_status_label") or "생성 대기").strip()
    return snapshot


def save_doc_job_snapshot(session, payload):
    snapshot = _normalize_doc_job_snapshot(payload)
    if not snapshot["job_id"]:
        return None
    snapshots = get_doc_job_snapshots(session)
    snapshots[snapshot["job_id"]] = snapshot
    session[DOC_JOB_SNAPSHOT_SESSION_KEY] = snapshots
    session.modified = True
    return snapshot


def _doc_job_snapshot_matches(snapshot, job_kind, document_code, *, job_id=None, tracking_document_sn=None):
    if not isinstance(snapshot, dict):
        return False
    if str(snapshot.get("job_kind") or "").strip() != str(job_kind or "").strip():
        return False
    if str(snapshot.get("docs_cd") or "").strip() != str(document_code or "").strip():
        return False
    if job_id and str(snapshot.get("job_id") or "").strip() != str(job_id).strip():
        return False
    if tracking_document_sn is not None and str(snapshot.get("tracking_document_sn") or "").strip() != str(tracking_document_sn).strip():
        return False
    return True


def find_doc_job_snapshot(session, job_kind, document_code, *, job_id=None, tracking_document_sn=None):
    snapshots = get_doc_job_snapshots(session)
    if job_id:
        snapshot = snapshots.get(str(job_id).strip())
        if _doc_job_snapshot_matches(snapshot, job_kind, document_code, job_id=job_id, tracking_document_sn=tracking_document_sn):
            return snapshot
    for snapshot in snapshots.values():
        if _doc_job_snapshot_matches(snapshot, job_kind, document_code, tracking_document_sn=tracking_document_sn):
            return snapshot
    return None


def clear_doc_job_snapshot(session, *, job_id=None, job_kind=None, document_code=None, tracking_document_sn=None):
    snapshots = get_doc_job_snapshots(session)
    if not snapshots:
        return False

    removed = False
    if job_id:
        removed = snapshots.pop(str(job_id).strip(), None) is not None
    else:
        remaining = {}
        for snapshot_job_id, snapshot in snapshots.items():
            if job_kind and document_code and _doc_job_snapshot_matches(
                snapshot,
                job_kind,
                document_code,
                tracking_document_sn=tracking_document_sn,
            ):
                removed = True
                continue
            remaining[snapshot_job_id] = snapshot
        snapshots = remaining

    if removed:
        if snapshots:
            session[DOC_JOB_SNAPSHOT_SESSION_KEY] = snapshots
        else:
            session.pop(DOC_JOB_SNAPSHOT_SESSION_KEY, None)
        session.modified = True
    return removed


def _normalize_reference_filename(filename):
    return Path(filename or "").name or "reference"


def build_itf_reference_storage_key(project, actor, filename):
    project_key = getattr(project, "sn", "none")
    actor_key = getattr(actor, "sn", "anonymous")
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else "bin"
    return f"temp/itf/{project_key}/{actor_key}/{uuid4().hex}.{extension}"


def _legacy_cleanup_path(path_value):
    if not path_value:
        return
    path = Path(path_value)
    for _ in range(3):
        try:
            path.unlink(missing_ok=True)
            return
        except FileNotFoundError:
            return
        except PermissionError:
            try:
                os.chmod(path, 0o666)
            except OSError:
                pass
            time.sleep(0.05)
        except OSError:
            return


def cleanup_generation_itf_reference(reference):
    storage_key = (reference or {}).get("storage_key", "")
    if storage_key:
        delete_object(storage_key)
        return
    _legacy_cleanup_path((reference or {}).get("path", ""))


def cleanup_generation_itf_references(state):
    for reference in state.get("itf_reference_files", []):
        cleanup_generation_itf_reference(reference)


def get_generation_itf_references(state):
    return list(state.get("itf_reference_files", []))


def get_fastapi_base_url():
    return str(getattr(settings, "FASTAPI_BASE_URL", "") or "").rstrip("/")


def get_fastapi_api_key():
    return str(getattr(settings, "FASTAPI_API_KEY", "") or "").strip()


def build_fastapi_json_headers():
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "User-Agent": FASTAPI_CLIENT_USER_AGENT,
    }
    api_key = get_fastapi_api_key()
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    return headers


def post_fastapi_json(url, payload, *, timeout_seconds):
    headers = build_fastapi_json_headers()
    with requests.Session() as session:
        # Keep the previous behavior that ignored HTTP(S)_PROXY environment variables.
        session.trust_env = False
        response = session.post(url, json=payload, headers=headers, timeout=timeout_seconds)
        response.raise_for_status()
        return response


def get_doc_job_poll_interval_seconds():
    raw_value = getattr(settings, "DOC_JOB_POLL_INTERVAL_SECONDS", 10)
    try:
        interval = int(raw_value)
    except (TypeError, ValueError):
        return 10
    return interval if interval > 0 else 10


def _safe_parse_positive_int(value):
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return None
    return parsed if parsed > 0 else None


def _build_reference_uri(reference):
    path_value = str((reference or {}).get("path", "") or "").strip()
    if path_value:
        return path_value

    storage_key = str((reference or {}).get("storage_key", "") or "").strip()
    if not storage_key:
        return ""

    try:
        return build_s3_uri(storage_key)
    except Exception:
        return storage_key


def add_generation_itf_references(project, actor, state, uploaded_files):
    references = state.setdefault("itf_reference_files", [])
    added_count = 0
    errors = []

    for uploaded_file in uploaded_files:
        if uploaded_file is None or not getattr(uploaded_file, "name", ""):
            continue

        filename = _normalize_reference_filename(uploaded_file.name)
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if extension not in INTERFACE_REFERENCE_ALLOWED_EXTENSIONS:
            errors.append("png, jpg, jpeg ?대?吏 ?뚯씪留??낅줈?쒗븷 ???덉뒿?덈떎.")
            continue
        if uploaded_file.size > INTERFACE_REFERENCE_MAX_FILE_SIZE:
            errors.append("媛??대?吏??3MB ?댄븯留??낅줈?쒗븷 ???덉뒿?덈떎.")
            continue

        token = uuid4().hex
        storage_key = build_itf_reference_storage_key(project, actor, filename)
        content_bytes = b"".join(uploaded_file.chunks())
        content_type = f"image/{'jpeg' if extension == 'jpg' else extension}"
        save_bytes(storage_key, content_bytes, content_type=content_type)

        references.append(
            {
                "token": token,
                "name": filename,
                "size": uploaded_file.size,
                "extension": extension,
                "storage_key": storage_key,
                "path": _build_reference_uri({"storage_key": storage_key}),
            }
        )
        added_count += 1

    state["itf_reference_files"] = references
    return added_count, errors


def remove_generation_itf_reference(state, token):
    if not token:
        return False

    remaining_references = []
    removed = False
    for reference in state.get("itf_reference_files", []):
        if reference.get("token") == token:
            cleanup_generation_itf_reference(reference)
            removed = True
            continue
        remaining_references.append(reference)

    state["itf_reference_files"] = remaining_references
    return removed


def _get_ordered_codes(code_values):
    code_map = Code.objects.in_bulk(code_values, field_name="code")
    return [code_map[code] for code in code_values if code in code_map]


def get_document_type_rows():
    return _get_ordered_codes(DOCUMENT_CODE_SEQUENCE)


def get_document_code_sequence():
    rows = get_document_type_rows()
    if rows:
        return [row.code for row in rows]
    return list(DOCUMENT_CODE_SEQUENCE)


def get_document_type_map():
    return {row.code: row for row in get_document_type_rows()}


def resolve_document_code(raw_code):
    sequence = get_document_code_sequence()
    if not sequence:
        return raw_code or DOCUMENT_CODE_SEQUENCE[0]
    return raw_code if raw_code in sequence else sequence[0]


def get_document_type_choices(*, include_all=False):
    rows = get_document_type_rows()
    choices = [("all", "전체 RA 문서")] if include_all else []
    for row in rows:
        choices.append((row.code, row.name))
    return tuple(choices)


def get_approval_status_choices(*, include_all=False):
    rows = _get_ordered_codes(APPROVAL_STATUS_SEQUENCE)
    choices = [("all", "전체 RA 문서")] if include_all else []
    for row in rows:
        choices.append((row.code, row.name))
    return tuple(choices)


def get_document_label(document_code):
    fallback_labels = {
        "DOC_SRS": "요구사항정의서",
        "DOC_ITF": "위험관리표",
        "DOC_ARCH": "소프트웨어 요구사항 명세서",
        "DOC_ERD": "추적성 매트릭스",
        "DOC_DB": "변경관리 기록",
        "DOC_TS": "통합시험 시나리오",
    }
    row = Code.objects.filter(code=document_code).only("name").first()
    return row.name if row else fallback_labels.get(document_code, document_code)


def get_document_index(document_code):
    sequence = get_document_code_sequence()
    try:
        return sequence.index(document_code)
    except ValueError:
        return 0


def get_previous_document_code(document_code):
    sequence = get_document_code_sequence()
    try:
        index = sequence.index(document_code)
    except ValueError:
        return None
    return None if index == 0 else sequence[index - 1]


def get_document_codes_before(document_code):
    sequence = get_document_code_sequence()
    try:
        index = sequence.index(document_code)
    except ValueError:
        return []
    return sequence[:index]


def get_document_prerequisite_codes(document_code):
    return tuple(DOCUMENT_PREREQUISITES.get(document_code, ()))


def get_generation_dependent_codes(document_code):
    target_code = resolve_document_code(document_code)
    if not target_code:
        return ()

    dependent_codes = set()
    pending_codes = [target_code]
    while pending_codes:
        current_code = pending_codes.pop()
        for candidate_code, prerequisite_codes in DOCUMENT_PREREQUISITES.items():
            if candidate_code == target_code or candidate_code in dependent_codes:
                continue
            if current_code in prerequisite_codes:
                dependent_codes.add(candidate_code)
                pending_codes.append(candidate_code)

    return tuple(code for code in get_document_code_sequence() if code in dependent_codes)


def get_generation_regeneration_target_codes(document_code):
    target_code = resolve_document_code(document_code)
    if not target_code:
        return ()
    return (target_code, *get_generation_dependent_codes(target_code))


def get_actor(request):
    return get_request_user(request)


def get_project_role(project, user):
    if project is None or user is None:
        return None
    role = (
        ProjectUserRole.objects.filter(
            project=project,
            user=user,
            role_id__in=PROJECT_ROLE_CODES,
        )
        .order_by("-role_id")
        .first()
    )
    return role.role_id if role else None


def is_project_manager(project, user):
    if getattr(user, "is_staff", False):
        return True
    return get_project_role(project, user) == "ROLE_MANAGER"


def is_project_participant(project, user):
    if getattr(user, "is_staff", False):
        return True
    return get_project_role(project, user) in PROJECT_ROLE_CODES


def get_document_title(document):
    return f"{document.document_type_id}_v{document.version}.docx"


def build_document_key(document, latest_detail=None):
    if latest_detail is None:
        latest_detail = get_latest_detail(document)
    detail_part = getattr(latest_detail, "sn", None) or "none"
    return f"docs-{document.sn}-v{document.version}-detail-{detail_part}"


def build_docx_bytes(title, body_lines):
    content = DocxDocument()
    content.add_heading(title, level=0)
    for line in body_lines:
        content.add_paragraph(line)
    buffer = io.BytesIO()
    content.save(buffer)
    return buffer.getvalue()


def extract_text_from_docx(binary_content):
    if not binary_content:
        return ""

    try:
        document = DocxDocument(io.BytesIO(binary_content))
    except Exception:
        return ""

    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text).strip()


def build_document_detail_storage_key(project, document_sn, detail_sn):
    return f"document-details/{project.sn}/{document_sn}/{detail_sn}.docx"


def build_document_detail_path(project, document_sn, detail_sn):
    return build_s3_uri(build_document_detail_storage_key(project, document_sn, detail_sn))


def get_document_detail_bytes(detail):
    if detail is None:
        return b""
    detail_path = str(getattr(detail, "path", "") or "").strip()
    if not detail_path:
        raise ValueError("Document detail is missing docs_path.")
    return read_bytes_from_uri(detail_path)


def get_docx_revision_fingerprint(content_bytes):
    content = content_bytes or b""
    try:
        with ZipFile(io.BytesIO(content)) as archive:
            document_xml = archive.read("word/document.xml")
    except (BadZipFile, KeyError):
        document_xml = content
    return hashlib.sha256(document_xml).hexdigest()


def is_same_docx_revision(left_bytes, right_bytes):
    return get_docx_revision_fingerprint(left_bytes) == get_docx_revision_fingerprint(right_bytes)


def download_remote_content(url):
    if not url:
        return None
    with urlopen(url, timeout=10) as response:
        return response.read()


def get_public_base_url(request=None):
    if request is not None:
        return request.build_absolute_uri("/").rstrip("/")
    return getattr(settings, "DJANGO_PUBLIC_BASE_URL", "").rstrip("/")


def normalize_onlyoffice_document_server_url(configured_url):
    normalized_url = (configured_url or "").strip().rstrip("/")
    api_suffix = "/web-apps/apps/api/documents/api.js"
    if normalized_url.endswith(api_suffix):
        normalized_url = normalized_url[: -len(api_suffix)]
    return normalized_url.rstrip("/")


def get_onlyoffice_document_server_url(request=None, *, browser=False):
    configured_url = normalize_onlyoffice_document_server_url(
        getattr(settings, "ONLYOFFICE_DOCUMENT_SERVER_URL", "")
    )
    if not configured_url:
        return ""

    if configured_url.startswith(("http://", "https://")):
        if browser:
            return configured_url
        return configured_url

    normalized_path = f"/{configured_url.lstrip('/')}".rstrip("/")
    if browser:
        return normalized_path or "/onlyoffice"

    public_base_url = get_public_base_url(request)
    if not public_base_url:
        raise ValueError("OnlyOffice Document Server URL requires a public base URL.")
    return urljoin(f"{public_base_url}/", normalized_path.lstrip("/")).rstrip("/")


def request_force_save(document, *, latest_detail=None, userdata=None, request=None):
    document_server_url = get_onlyoffice_document_server_url(request)
    if not document_server_url:
        raise ValueError("OnlyOffice Document Server URL is not configured.")

    document_key = build_document_key(document, latest_detail=latest_detail)
    command_url = f"{document_server_url}/command?shardkey={document_key}"
    payload = {
        "c": "forcesave",
        "key": document_key,
    }
    if userdata:
        payload["userdata"] = userdata

    if settings.ONLYOFFICE_JWT_SECRET:
        request_payload = {
            "token": encode_jwt(payload, settings.ONLYOFFICE_JWT_SECRET),
        }
    else:
        request_payload = payload

    request = Request(
        command_url,
        data=json.dumps(request_payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urlopen(request, timeout=10) as response:
        return json.loads(response.read().decode("utf-8") or "{}")


def get_latest_detail(document):
    return document.details.filter(is_deleted="N").order_by("-created_at", "-sn").first()


def get_highest_detail_sn(document):
    if document is None:
        return None
    return document.details.filter(is_deleted="N").order_by("-sn").first()


def get_detail_by_sn(document, detail_sn):
    return document.details.filter(sn=detail_sn, is_deleted="N").order_by("-created_at", "-sn").first()


def get_latest_pending_approval(document):
    return (
        DocumentApproval.objects.filter(
            detail__document=document,
            approval_status_id="APRV_REQ",
        )
        .select_related("detail", "created_by", "approval_status")
        .order_by("-created_at", "-approval_sn")
        .first()
    )


def latest_confirmed_document(project, document_code, *, exclude_document_sn=None):
    queryset = (
        _exclude_working_versions(Document.objects.filter(project=project, document_type_id=document_code))
        .select_related("document_type", "created_by", "possession_user")
    )
    if exclude_document_sn is not None:
        queryset = queryset.exclude(sn=exclude_document_sn)
    return queryset.order_by("-created_at", "-sn").first()


def _exclude_working_versions(queryset):
    return queryset.exclude(version__in=WORKING_DOCUMENT_VERSIONS)


def _filter_working_versions(queryset):
    return queryset.filter(version__in=WORKING_DOCUMENT_VERSIONS)


def is_working_document(document):
    return str(getattr(document, "version", "") or "") in WORKING_DOCUMENT_VERSIONS


def _generation_completed_queryset(project, document_code):
    return (
        Document.objects.filter(
            project=project,
            document_type_id=document_code,
            progress_status_id=PROGRESS_COMPLETED,
        )
        .exclude(version="0.0")
        .select_related("project", "document_type", "created_by", "updated_by", "possession_user")
        .order_by("-updated_at", "-created_at", "-sn")
    )


def get_generation_saved_document(project, document_code, state=None):
    if project is None:
        return None

    confirmed_document = max(
        _generation_completed_queryset(project, document_code),
        key=lambda document: (
            document.sn,
        ),
        default=None,
    )
    if state:
        document_sn = (state.get("confirmed_documents", {}) or {}).get(document_code) or (
            state.get("confirmed_documents", {}) or {}
        ).get(str(document_code))
        if document_sn:
            document = (
                Document.objects.filter(
                    sn=document_sn,
                    project=project,
                    document_type_id=document_code,
                )
                .select_related("project", "document_type", "created_by", "updated_by", "possession_user")
                .first()
            )
            if document is not None:
                if document.progress_status_id == PROGRESS_COMPLETED and confirmed_document is None:
                    return document
                if confirmed_document is not None:
                    return confirmed_document

    return confirmed_document


def get_document_history_queryset(project, document_code=None):
    if project is None:
        return Document.objects.none()

    queryset = _exclude_working_versions(Document.objects.filter(project=project))
    if document_code and document_code != "all":
        queryset = queryset.filter(document_type_id=document_code)

    return (
        queryset.annotate(
            version_rank=Window(
                expression=RowNumber(),
                partition_by=[F("document_type_id"), F("version")],
                order_by=[F("created_at").desc(), F("sn").desc()],
            )
        )
        .filter(version_rank=1)
        .select_related("document_type", "created_by", "possession_user")
        .order_by("document_type_id", "-created_at", "-sn")
    )


def has_any_confirmed_initial_document(project):
    if project is None:
        return False
    return _exclude_working_versions(
        Document.objects.filter(
            project=project,
            document_type_id__in=get_document_code_sequence(),
        )
    ).exists()


def has_all_generated_document_types(project):
    if project is None:
        return False
    generated_count = (
        _exclude_working_versions(
            Document.objects.filter(
                project=project,
                document_type_id__in=get_document_code_sequence(),
            )
        )
        .values_list("document_type_id", flat=True)
        .distinct()
        .count()
    )
    return generated_count >= len(get_document_code_sequence())


def can_start_initial_generation(project, user):
    if project is None or user is None or not is_project_manager(project, user):
        return False
    return not has_all_generated_document_types(project)


def has_active_generation_session(state):
    return (
        bool(
            state.get("selected_file_ids")
            or state.get("draft_documents")
            or state.get("confirmed_documents")
            or state.get("itf_reference_files")
        )
        and not is_generation_complete(state)
    )


def can_access_initial_generation(project, user, state):
    """Return whether the user can start or continue the sequential document generation flow.

    Project managers and assigned project members can generate deliverables.
    If a generation session is already active, project participants can continue it.
    If all deliverable types already have confirmed versions, generation is blocked until the
    user explicitly starts a regeneration flow.
    """
    if project is None or user is None or not is_project_participant(project, user):
        return False
    return has_active_generation_session(state) or not has_all_generated_document_types(project)


def is_latest_document_for_type(document):
    if document is None:
        return False
    latest_document = (
        Document.objects.filter(
            project=document.project,
            document_type_id=document.document_type_id,
        )
        .order_by("-sn")
        .first()
    )
    return latest_document is not None and latest_document.sn == document.sn


def is_latest_detail_for_document(document, detail):
    if document is None or detail is None:
        return False
    latest_detail = get_highest_detail_sn(document)
    return latest_detail is not None and latest_detail.sn == detail.sn


def _build_fastapi_generate_url():
    base_url = get_fastapi_base_url()
    if not base_url:
        raise ValueError("FastAPI base URL is not configured.")
    return f"{base_url}/generate"


def request_fastapi_generate(payload):
    url = _build_fastapi_generate_url()
    timeout_seconds = FASTAPI_GENERATE_TIMEOUT_SECONDS
    headers = build_fastapi_json_headers()
    _debug_generation_log(
        "fastapi_request_prepare",
        url=url,
        timeout_seconds=timeout_seconds,
        authorization_present="Authorization" in headers,
        user_agent=headers.get("User-Agent"),
        payload=_summarize_generation_payload(payload),
    )
    try:
        response = post_fastapi_json(url, payload, timeout_seconds=timeout_seconds)
        status_code = response.status_code
        body = response.text or "{}"
    except Exception as exc:
        response = getattr(exc, "response", None)
        _debug_generation_log(
            "fastapi_request_error",
            url=url,
            error_type=type(exc).__name__,
            error=str(exc),
            status_code=getattr(response, "status_code", None),
            body_preview=_truncate_log_value(getattr(response, "text", "") or ""),
        )
        traceback.print_exc()
        raise
    _debug_generation_log(
        "fastapi_response_received",
        url=url,
        status_code=status_code,
        body_preview=_truncate_log_value(body),
    )
    try:
        return json.loads(body)
    except json.JSONDecodeError:
        _debug_generation_log("fastapi_response_non_json", url=url)
        return {"raw": body}


def request_fastapi_approval_review(approval_sn):
    base_url = get_fastapi_base_url()
    if not base_url:
        raise ValueError("FastAPI base URL is not configured.")

    url = f"{base_url}/approval-review"
    payload = {"docs_aprv_sn": approval_sn}
    response = post_fastapi_json(url, payload, timeout_seconds=FASTAPI_APPROVAL_REVIEW_TIMEOUT_SECONDS)
    body = response.text or "{}"
    try:
        return json.loads(body)
    except json.JSONDecodeError:
        return {"raw": body}


def extract_fastapi_error_message(exc):
    if isinstance(exc, requests.HTTPError):
        response = getattr(exc, "response", None)
        status_code = getattr(response, "status_code", None)
        body = getattr(response, "text", "") or ""
        if body:
            try:
                payload = json.loads(body)
            except json.JSONDecodeError:
                payload = None
            if isinstance(payload, dict):
                detail = payload.get("detail")
                if isinstance(detail, str) and detail.strip():
                    return f"FastAPI error {status_code}: {detail.strip()}"
                if isinstance(detail, list) and detail:
                    first_detail = detail[0]
                    if isinstance(first_detail, dict):
                        message = first_detail.get("msg") or first_detail.get("message")
                        if message:
                            return f"FastAPI error {status_code}: {message}"
        return f"FastAPI error {status_code}: {exc}"
    if isinstance(exc, requests.RequestException):
        return f"FastAPI connection failed: {exc}"
    if isinstance(exc, HTTPError):
        try:
            body = exc.read().decode("utf-8")
        except Exception:
            body = ""
        if body:
            try:
                payload = json.loads(body)
            except json.JSONDecodeError:
                payload = None
            if isinstance(payload, dict):
                detail = payload.get("detail")
                if isinstance(detail, str) and detail.strip():
                    return f"FastAPI error {exc.code}: {detail.strip()}"
                if isinstance(detail, list) and detail:
                    first_detail = detail[0]
                    if isinstance(first_detail, dict):
                        message = first_detail.get("msg") or first_detail.get("message")
                        if message:
                            return f"FastAPI error {exc.code}: {message}"
        return f"FastAPI error {exc.code}: {exc.reason}"
    if isinstance(exc, URLError):
        reason = getattr(exc, "reason", None)
        return f"FastAPI connection failed: {reason or exc}"
    if isinstance(exc, ValueError):
        return str(exc)
    return str(exc)


def get_generation_reference_uris(state):
    uris = []
    for reference in get_generation_itf_references(state):
        uri = _build_reference_uri(reference)
        if uri:
            uris.append(uri)
    return uris

#### fast api ?붿껌 payload 留뚮뱾湲?def build_generation_request_payload(project, state, document_code, *, update_mode="N", selected_files=None):
    image_list = get_generation_reference_uris(state) if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE else []
    if document_code in FILE_INPUT_DOCUMENT_CODES:
        files = selected_files if selected_files is not None else get_generation_selected_files(project, state)
    else:
        # DOC_ITF??image_list, DOC_ARCH??tbl_project_net, DOC_ERD/DOC_DB/DOC_TS??        # project_sn?쇰줈 ?꾩슂???좏뻾 ?곗텧臾쇱쓣 FastAPI媛 議고쉶?섎뒗 援ъ“?낅땲??
        files = []
    return {
        "project_sn": project.sn,
        "docs_cd": document_code,
        "udt_yn": update_mode,
        "file_list": [project_file.sn for project_file in files],
        "image_list": image_list,
        "etc": {"debug": bool(getattr(settings, "DEBUG", False))},
    }


def build_auto_apply_request_payload(project, document_code, selected_files):
    return {
        "project_sn": project.sn,
        "docs_cd": document_code,
        "udt_yn": "Y",
        "file_list": [project_file.sn for project_file in selected_files],
        "image_list": [],
        "etc": {"debug": bool(getattr(settings, "DEBUG", False))},
    }


def get_generation_job_kind(job):
    if job is None:
        return GENERATION_JOB_KIND_INITIAL
    request_payload = getattr(job, "request_payload", None) or {}
    update_mode = str(request_payload.get("udt_yn", "N") or "N").upper()
    if update_mode == "Y":
        return GENERATION_JOB_KIND_AUTO_APPLY
    return GENERATION_JOB_KIND_INITIAL


def get_generation_job_label(job):
    return get_document_label(getattr(job, "document_type_id", None))


def parse_fastapi_generation_response(payload, *, fallback_document_code=None):
    if not isinstance(payload, dict):
        return {}
    raw_job_id = payload.get("job_id")
    job_id = str(raw_job_id).strip() if raw_job_id is not None else ""
    if not job_id:
        return {}
    raw_status = str(payload.get("status") or "").strip()
    job_status_code = str(payload.get("job_status_code") or "").strip()
    if not job_status_code:
        if raw_status in {"started", "accepted", "queued", "pending"}:
            job_status_code = PROGRESS_PENDING
        elif raw_status == "running":
            job_status_code = PROGRESS_PROCESSING
        elif raw_status == "failed":
            job_status_code = PROGRESS_FAILED
        elif raw_status == "completed":
            job_status_code = PROGRESS_COMPLETED
        else:
            job_status_code = PROGRESS_PENDING
    return {
        "job_id": job_id,
        "request_id": str(payload.get("request_id") or "").strip(),
        "project_sn": _safe_parse_positive_int(payload.get("project_sn")),
        "docs_cd": str(payload.get("docs_cd") or fallback_document_code or "").strip(),
        "status": raw_status or "started",
        "status_url": str(payload.get("status_url") or "").strip(),
        "job_status_code": job_status_code,
        "job_status_label": str(payload.get("job_status_label") or "").strip(),
        "message": str(payload.get("message") or "").strip() or "Generation accepted.",
    }


def _generation_job_queryset(project, document_code=None):
    queryset = GenerationJob.objects.select_related(
        "project",
        "document_type",
        "document",
        "job_status",
    )
    if project is not None:
        queryset = queryset.filter(project=project)
    if document_code:
        queryset = queryset.filter(document_type_id=document_code)
    return queryset


def _matches_generation_job_kind(job, job_kind=None):
    if not job_kind:
        return True
    return get_generation_job_kind(job) == job_kind


def _filter_generation_job_kind(queryset, job_kind=None):
    if job_kind == GENERATION_JOB_KIND_AUTO_APPLY:
        return queryset.filter(request_payload__udt_yn="Y")
    if job_kind == GENERATION_JOB_KIND_INITIAL:
        return queryset.exclude(request_payload__udt_yn="Y")
    return queryset


def get_generation_job(project, *, job_sn=None, job_id=None):
    queryset = _generation_job_queryset(project)
    if job_sn is not None:
        return queryset.filter(sn=job_sn).first()
    if job_id:
        return queryset.filter(job_id=str(job_id).strip()).first()
    return None


def find_generation_job(
    project,
    document_code,
    *,
    job_kind=None,
    job_id=None,
    tracking_document_sn=None,
    status_codes=None,
):
    queryset = _generation_job_queryset(project, document_code)
    if status_codes:
        queryset = queryset.filter(job_status_id__in=status_codes)
    queryset = _filter_generation_job_kind(queryset, job_kind)
    if job_id:
        return queryset.filter(job_id=str(job_id).strip()).first()
    if tracking_document_sn:
        tracked_job = queryset.filter(document_id=tracking_document_sn).order_by("-sn").first()
        if tracked_job is not None:
            return tracked_job
    return queryset.order_by("-sn").first()


def get_running_generation_job(project, document_code, *, job_kind=None, tracking_document_sn=None):
    return find_generation_job(
        project,
        document_code,
        job_kind=job_kind,
        tracking_document_sn=tracking_document_sn,
        status_codes=RUNNING_PROGRESS_CODES,
    )


def get_latest_generation_job(project, document_code, *, job_kind=None):
    return find_generation_job(project, document_code, job_kind=job_kind)


def get_running_initial_job(project, document_code, *, tracking_document_sn=None):
    return get_running_generation_job(
        project,
        document_code,
        job_kind=GENERATION_JOB_KIND_INITIAL,
        tracking_document_sn=tracking_document_sn,
    )


def get_running_auto_apply_job(project, document_code, *, tracking_document_sn=None):
    return get_running_generation_job(
        project,
        document_code,
        job_kind=GENERATION_JOB_KIND_AUTO_APPLY,
        tracking_document_sn=tracking_document_sn,
    )


def get_running_history_job(project, document_code):
    job = get_running_generation_job(project, document_code)
    if job is None:
        return None, None
    return get_generation_job_kind(job), job


def _document_job_queryset(project, document_code, *, initial_only=False):
    queryset = Document.objects.filter(project=project, document_type_id=document_code).select_related(
        "project",
        "document_type",
        "created_by",
        "updated_by",
        "possession_user",
    )
    if initial_only:
        queryset = queryset.filter(version__in=WORKING_DOCUMENT_VERSIONS)
    return queryset


def get_generation_draft_document(project, state, document_code=None):
    if project is None:
        return None
    target_code = document_code or get_current_generation_code(state, project)
    if not target_code:
        return None
    target_sn = state.get("draft_documents", {}).get(target_code)
    if not target_sn:
        return None
    return (
        _document_job_queryset(project, target_code, initial_only=True)
        .filter(sn=target_sn)
        .first()
    )


def set_generation_draft_document(state, document):
    if document is None:
        return state
    state.setdefault("draft_documents", {})[document.document_type_id] = document.sn
    return state


def set_generation_document_from_job(state, document):
    if document is None:
        return state
    code = document.document_type_id
    if is_working_document(document):
        state.setdefault("draft_documents", {})[code] = document.sn
        state.setdefault("confirmed_documents", {}).pop(code, None)
        state.setdefault("confirmed_documents", {}).pop(str(code), None)
    else:
        state.setdefault("confirmed_documents", {})[code] = document.sn
        state.setdefault("draft_documents", {}).pop(code, None)
        state.setdefault("draft_documents", {}).pop(str(code), None)
    return state


def clear_generation_draft_document(state, document_code):
    state.setdefault("draft_documents", {}).pop(document_code, None)
    return state


def start_initial_generation_job(project, actor, state, document_code=None):
    document_code = get_current_generation_code(state, project, preferred_code=document_code)
    _debug_generation_log(
        "start_initial_generation_job_enter",
        project_sn=getattr(project, "sn", None),
        actor_sn=getattr(actor, "sn", None),
        document_code=document_code,
        selected_file_ids=list(state.get("selected_file_ids", []) or []),
        draft_documents=dict(state.get("draft_documents", {}) or {}),
        confirmed_documents=dict(state.get("confirmed_documents", {}) or {}),
        itf_reference_count=len(state.get("itf_reference_files", []) or []),
    )
    if not document_code:
        _debug_generation_log("start_initial_generation_job_no_document_code")
        return {"status": "error", "job": None, "document": None, "message": "?앹꽦???곗텧臾??④퀎瑜?李얠? 紐삵뻽?듬땲??"}

    running_job = get_running_initial_job(project, document_code)
    if running_job is not None:
        _debug_generation_log(
            "start_initial_generation_job_already_running",
            job_id=running_job.job_id,
            job_status_id=running_job.job_status_id,
            document_sn=running_job.document_id,
        )
        if running_job.document is not None:
            set_generation_draft_document(state, running_job.document)
        return {
            "status": "running",
            "job": running_job,
            "document": running_job.document,
            "message": "臾몄꽌瑜??앹꽦 以묒엯?덈떎.",
        }

    payload = build_generation_request_payload(project, state, document_code, update_mode="N")
    _debug_generation_log(
        "start_initial_generation_job_payload_built",
        payload=_summarize_generation_payload(payload),
    )
    try:
        fastapi_response = request_fastapi_generate(payload)
        _debug_generation_log(
            "start_initial_generation_job_fastapi_accepted",
            response=_truncate_log_value(fastapi_response),
        )
    except (HTTPError, URLError, ValueError) as exc:
        _debug_generation_log(
            "start_initial_generation_job_fastapi_failed",
            error_type=type(exc).__name__,
            error=extract_fastapi_error_message(exc),
        )
        return {"status": "error", "job": None, "document": None, "message": extract_fastapi_error_message(exc)}

    accepted_response = parse_fastapi_generation_response(fastapi_response, fallback_document_code=document_code)
    if not accepted_response.get("job_id"):
        _debug_generation_log("start_initial_generation_job_missing_job_id", response=_truncate_log_value(fastapi_response))
        return {"status": "error", "job": None, "document": None, "message": "臾몄꽌 ?앹꽦 ?묒뾽???쒖옉?섏? 紐삵뻽?듬땲??"}

    tracked_job = get_generation_job(project, job_id=accepted_response["job_id"])
    tracked_document = tracked_job.document if tracked_job is not None else None
    if tracked_document is not None:
        set_generation_draft_document(state, tracked_document)
    _debug_generation_log(
        "start_initial_generation_job_started",
        job_id=accepted_response["job_id"],
        tracked_document_sn=getattr(tracked_document, "sn", None),
        job_status_id=getattr(tracked_job, "job_status_id", None),
        version=getattr(tracked_document, "version", None),
    )
    return {
        "status": "started",
        "job": tracked_job or SimpleNamespace(**accepted_response),
        "document": tracked_document,
        "message": accepted_response.get("message") or "臾몄꽌 ?앹꽦???붿껌?덉뒿?덈떎.",
    }


def start_auto_apply_job(project, document_code, selected_files):
    _debug_generation_log(
        "start_auto_apply_job_enter",
        project_sn=getattr(project, "sn", None),
        document_code=document_code,
        selected_file_ids=[project_file.sn for project_file in selected_files],
    )
    running_job = get_running_auto_apply_job(project, document_code)
    if running_job is not None:
        _debug_generation_log(
            "start_auto_apply_job_already_running",
            job_id=running_job.job_id,
            job_status_id=running_job.job_status_id,
            document_sn=running_job.document_id,
        )
        return {
            "status": "running",
            "job": running_job,
            "document": running_job.document,
            "message": "臾몄꽌瑜??앹꽦 以묒엯?덈떎.",
        }

    payload = build_auto_apply_request_payload(project, document_code, selected_files)
    _debug_generation_log(
        "start_auto_apply_job_payload_built",
        payload=_summarize_generation_payload(payload),
    )
    try:
        fastapi_response = request_fastapi_generate(payload)
        _debug_generation_log(
            "start_auto_apply_job_fastapi_accepted",
            response=_truncate_log_value(fastapi_response),
        )
    except (HTTPError, URLError, ValueError) as exc:
        _debug_generation_log(
            "start_auto_apply_job_fastapi_failed",
            error_type=type(exc).__name__,
            error=extract_fastapi_error_message(exc),
        )
        return {"status": "error", "job": None, "document": None, "message": extract_fastapi_error_message(exc)}

    accepted_response = parse_fastapi_generation_response(fastapi_response, fallback_document_code=document_code)
    if not accepted_response.get("job_id"):
        _debug_generation_log("start_auto_apply_job_missing_job_id", response=_truncate_log_value(fastapi_response))
        return {"status": "error", "job": None, "document": None, "message": "?뚯쓽 ?댁슜 ?먮룞 ?곸슜 ?묒뾽???쒖옉?섏? 紐삵뻽?듬땲??"}

    tracked_job = get_generation_job(project, job_id=accepted_response["job_id"])
    tracked_document = tracked_job.document if tracked_job is not None else None
    _debug_generation_log(
        "start_auto_apply_job_started",
        job_id=accepted_response["job_id"],
        tracked_document_sn=getattr(tracked_document, "sn", None),
        job_status_id=getattr(tracked_job, "job_status_id", None),
        version=getattr(tracked_document, "version", None),
    )
    return {
        "status": "started",
        "job": tracked_job or SimpleNamespace(**accepted_response),
        "document": tracked_document,
        "message": accepted_response.get("message") or "?뚯쓽 ?댁슜 ?먮룞 ?곸슜???붿껌?덉뒿?덈떎.",
    }


def build_generation_lines(project, document_code, inputs):
    label = get_document_label(document_code)
    project_name = getattr(project, "name", "RA 프로젝트")
    rows = [
        f"{project_name} 프로젝트의 {label} 초안입니다.",
        "본 초안은 의료 진단 또는 치료 판단을 수행하지 않으며, 의료기기 개발 및 RA 문서 작성 과정을 보조하기 위한 문서입니다.",
    ]

    if document_code == "DOC_SRS":
        rows.append("업로드된 의료기기 개발 자료를 기반으로 사용 목적, 대상 사용자, 기능/비기능 요구사항, 규제 요구사항을 구조화했습니다.")
        for project_file in inputs:
            file_type_name = getattr(getattr(project_file, "file_type", None), "name", "근거 자료")
            rows.append(f"- 근거 자료: {project_file.name} ({file_type_name})")
    elif document_code == "DOC_ITF":
        rows.append("요구사항정의서를 기준으로 위해요인, 위해상황, 위험통제 방법, 검증항목을 연결했습니다.")
        for reference in inputs:
            rows.append(f"- 참고 자료: {reference.get('name', '위험관리 참고 자료')}")
    elif document_code == "DOC_ARCH":
        rows.append("SaMD 개발문서 보조를 위해 소프트웨어 요구사항, 구성요소, 데이터 흐름, 보안 요구사항을 정리했습니다.")
        for project_net in inputs:
            description = project_net.purpose or "목적 미입력"
            rows.append(f"- 구성요소: {project_net.name} ({description})")
    elif document_code == "DOC_ERD":
        rows.append("요구사항, 위험, 시험, 검증 결과를 연결하는 추적성 매트릭스 초안을 구성했습니다.")
        for previous_document in inputs:
            rows.append(f"- 선행 문서: {get_document_label(previous_document.document_type_id)} v{previous_document.version}")
    elif document_code == "DOC_DB":
        rows.append("요구사항 또는 설계 변경이 위험관리, 검증, 인허가 영향에 미치는 항목을 변경관리 기록으로 정리했습니다.")
        for previous_document in inputs:
            rows.append(f"- 영향 분석 기준 문서: {get_document_label(previous_document.document_type_id)} v{previous_document.version}")
    elif document_code == "DOC_TS":
        rows.append("요구사항과 위험통제 항목을 기반으로 통합시험 시나리오와 합격 기준을 구성했습니다.")
        for previous_document in inputs:
            rows.append(f"- 시험 설계 기준 문서: {get_document_label(previous_document.document_type_id)} v{previous_document.version}")
    else:
        rows.append("선택한 입력 자료를 기반으로 RA 문서 초안을 구성했습니다.")

    prerequisite_codes = get_document_prerequisite_codes(document_code)
    if prerequisite_codes:
        prerequisite_labels = ", ".join(get_document_label(code) for code in prerequisite_codes)
        rows.append(f"이 문서는 선행 RA 문서인 {prerequisite_labels} 저장본을 기준으로 작성됩니다.")
    rows.append("생성된 초안은 담당자가 검토한 뒤 OnlyOffice 또는 상세 화면에서 보완하고 정합성 검토 요청을 진행해야 합니다.")
    return rows

@transaction.atomic
def create_document_with_detail(
    *,
    project,
    document_code,
    actor,
    version,
    modification_content,
    content_bytes,
    locked_user=None,
    progress_status_id=PROGRESS_COMPLETED,
):
    try:
        document = Document.objects.create(
            project=project,
            possession_user=locked_user,
            document_type_id=document_code,
            progress_status_id=progress_status_id,
            version=version,
            modification_content=modification_content,
            created_by=actor,
            updated_by=actor,
        )
        detail = DocumentDetail.objects.create(
            document=document,
            path="",
            is_deleted="N",
            created_by=actor,
        )
        detail_path = build_document_detail_path(project, document.sn, detail.sn)
        save_bytes(
            build_document_detail_storage_key(project, document.sn, detail.sn),
            content_bytes or b"",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        detail.path = detail_path
        detail.save(update_fields=["path"])
    except Exception:
        if "detail_path" in locals():
            delete_object_at_uri(detail_path)
        raise
    return document, detail


def create_draft_document(project, document_code, actor, source_inputs):
    label = get_document_label(document_code)
    content_bytes = build_docx_bytes(label, build_generation_lines(project, document_code, source_inputs))
    return create_document_with_detail(
        project=project,
        document_code=document_code,
        actor=actor,
        version="0.0",
        modification_content="理쒖큹 ?앹꽦",
        content_bytes=content_bytes,
        locked_user=None,
        progress_status_id=PROGRESS_PROCESSING,
    )


def _parse_document_version_number(version):
    text = str(version or "").strip().lower()
    if not text or text == "0":
        return None
    if text.startswith("v"):
        text = text[1:]
    if "." in text:
        text = text.split(".", 1)[0]
    try:
        number = int(text)
    except (TypeError, ValueError):
        return None
    return number if number > 0 else None


def get_next_document_version(project, document_code):
    numbers = []
    versions = Document.objects.filter(
        project=project,
        document_type_id=document_code,
    )
    versions = _exclude_working_versions(versions).values_list("version", flat=True)
    for version in versions:
        number = _parse_document_version_number(version)
        if number is not None:
            numbers.append(number)
    return str((max(numbers) if numbers else 0) + 1)


@transaction.atomic
def confirm_document(document, actor):
    latest_detail = get_latest_detail(document)
    if latest_detail is None:
        raise ValueError("Document detail is required.")
    document.possession_user = None
    document.progress_status_id = PROGRESS_COMPLETED
    document.updated_by = actor
    document.modification_content = "RA 문서 확정"
    document.save(update_fields=["possession_user", "progress_status", "updated_by", "modification_content"])
    return document, latest_detail


@transaction.atomic
def acquire_document_lock(document, actor):
    if document.possession_user_id and document.possession_user_id != actor.sn:
        return False
    document.possession_user = actor
    document.updated_by = actor
    document.save(update_fields=["possession_user", "updated_by"])
    return True


@transaction.atomic
def release_document_lock(document, actor=None):
    document.possession_user = None
    if actor is not None:
        document.updated_by = actor
        document.save(update_fields=["possession_user", "updated_by"])
    else:
        document.save(update_fields=["possession_user"])


@transaction.atomic
def save_revision(document, actor, *, text_content=None, content_bytes=None, modification_content="수정 저장"):
    latest_detail = get_latest_detail(document)
    if content_bytes is None:
        if text_content is not None:
            content_bytes = build_docx_bytes(
                get_document_label(document.document_type_id),
                text_content.splitlines(),
            )
        elif latest_detail is not None:
            content_bytes = get_document_detail_bytes(latest_detail)

    try:
        detail = DocumentDetail.objects.create(
            document=document,
            path="",
            is_deleted="N",
            created_by=actor,
        )
        detail_path = build_document_detail_path(document.project, document.sn, detail.sn)
        save_bytes(
            build_document_detail_storage_key(document.project, document.sn, detail.sn),
            content_bytes or b"",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        detail.path = detail_path
        detail.save(update_fields=["path"])
    except Exception:
        if "detail_path" in locals():
            delete_object_at_uri(detail_path)
        raise
    document.updated_by = actor
    document.modification_content = "RA 문서 확정"
    document.save(update_fields=["updated_by", "modification_content"])
    return detail


def apply_meeting_notes(document, actor, selected_files):
    base_text = extract_text_from_docx(get_document_detail_bytes(get_latest_detail(document)))
    new_text = [
        base_text,
        "",
        "[?뚯쓽 ?댁슜 ?먮룞 諛섏쁺]",
        "?ㅼ쓬 ?뚯쓽濡앹쓣 諛섏쁺???붾? ?섏젙 寃곌낵?낅땲??",
    ]
    for project_file in selected_files:
        new_text.append(f"- {project_file.name}")
    return save_revision(
        document,
        actor,
        text_content="\n".join(new_text).strip(),
        modification_content="?뚯쓽 ?댁슜 ?먮룞 諛섏쁺",
    )


@transaction.atomic
def restore_revision(document, actor, source_detail):
    return save_revision(
        document,
        actor,
        content_bytes=get_document_detail_bytes(source_detail),
        modification_content="?댁쟾 踰꾩쟾 蹂듭썝",
    )


def can_request_approval(document, actor, *, pending_approval=None):
    if actor is None or pending_approval is not None:
        return False
    return document.updated_by_id == actor.sn


@transaction.atomic
def create_approval_request(document, actor, request_content):
    latest_detail = get_latest_detail(document)
    approval = DocumentApproval.objects.create(
        detail=latest_detail,
        approval_status_id="APRV_REQ",
        request_content=request_content,
        rejection_reason=None,
        created_by=actor,
        updated_by=actor,
    )
    release_document_lock(document, actor)
    return approval


def get_latest_approval_review_job(approval):
    if approval is None:
        return None
    latest_sn = ApprovalReviewJob.objects.filter(approval_id=approval.approval_sn).aggregate(
        latest_sn=Max("sn"),
    )["latest_sn"]
    if latest_sn is None:
        return None
    return (
        ApprovalReviewJob.objects.filter(sn=latest_sn)
        .select_related("before_detail", "after_detail", "approval_request_detail")
        .first()
    )


@transaction.atomic
def cancel_approval_request(approval):
    approval.delete()


def has_document_version(project, document_code, version):
    return Document.objects.filter(
        project=project,
        document_type_id=document_code,
        version=version,
    ).exists()


@transaction.atomic
def approve_request(approval, actor, new_version, modification_content=None):
    source_detail = approval.detail
    source_document = source_detail.document
    document, detail = create_document_with_detail(
        project=source_document.project,
        document_code=source_document.document_type_id,
        actor=actor,
        version=new_version,
        modification_content=modification_content or approval.request_content or "?뱀씤 諛섏쁺",
        content_bytes=get_document_detail_bytes(source_detail),
        locked_user=None,
    )
    approval.approval_status_id = "APRV_COM"
    approval.updated_by = actor
    approval.save(update_fields=["approval_status", "updated_by"])
    return document, detail


@transaction.atomic
def reject_request(approval, actor, reason):
    approval.approval_status_id = "APRV_RJT"
    approval.rejection_reason = reason
    approval.updated_by = actor
    approval.save(update_fields=["approval_status", "rejection_reason", "updated_by"])


def hydrate_generation_state_from_existing_documents(project, state):
    """
    釉뚮씪?곗? ?몄뀡??鍮꾩뼱 ?덉뼱??DB???대? ????꾨즺???곗텧臾쇱씠 ?덉쑝硫?    ?쒖감 ?앹꽦 吏꾪뻾 ?곹깭瑜??댁뼱諛쏆뒿?덈떎.

    ?? ?ъ깮??紐⑤뱶(regeneration_targets)媛 ?덉쑝硫??좏깮???곗텧臾쇰쭔
    ????꾨즺 ?곹깭 蹂듦뎄?먯꽌 ?쒖쇅?⑸땲?? ?덈? ?ㅼ뼱 DOC_DB ?ъ깮?깆씠硫?    DOC_DB留??앹꽦 ?湲곕줈 ?먭퀬 DOC_TS ???ㅻⅨ ????꾨즺 ?곗텧臾쇱? ?좎??⑸땲??
    """
    if project is None:
        return state

    confirmed = state.setdefault("confirmed_documents", {})
    draft_documents = state.setdefault("draft_documents", {})
    sequence = get_document_code_sequence()
    regeneration_targets = set(state.get("regeneration_targets") or ())
    legacy_regeneration_from = state.get("regeneration_from")
    if legacy_regeneration_from in sequence and not regeneration_targets:
        regeneration_targets.add(legacy_regeneration_from)
    restorable_regeneration_targets = {legacy_regeneration_from} if legacy_regeneration_from in sequence else set()
    regeneration_job_baseline_sns = state.get("regeneration_job_baseline_sns") or {}

    for code in sequence:
        if code in regeneration_targets:
            confirmed.pop(str(code), None)
            confirmed.pop(code, None)
            draft_documents.pop(str(code), None)
            draft_documents.pop(code, None)
            if code in restorable_regeneration_targets:
                latest_job = find_generation_job(project, code, job_kind=GENERATION_JOB_KIND_INITIAL)
                job_document = getattr(latest_job, "document", None)
                if (
                    latest_job is not None
                    and latest_job.job_status_id == PROGRESS_COMPLETED
                    and job_document is not None
                    and code in regeneration_job_baseline_sns
                    and latest_job.sn > int(regeneration_job_baseline_sns.get(code) or 0)
                ):
                    if is_working_document(job_document):
                        draft_documents[str(code)] = job_document.sn
                    else:
                        confirmed[str(code)] = job_document.sn
            continue
        if str(code) in confirmed:
            confirmed_document = get_generation_saved_document(project, code, state)
            if confirmed_document is not None:
                confirmed[str(code)] = confirmed_document.sn
            else:
                confirmed.pop(str(code), None)
                confirmed.pop(code, None)
            continue
        if str(code) in draft_documents:
            continue

        confirmed_document = get_generation_saved_document(project, code, state)
        if confirmed_document is None:
            continue
        confirmed[str(code)] = confirmed_document.sn

    return state


def begin_generation_regeneration(session, project, document_code):
    target_code = resolve_document_code(document_code)
    state = _build_empty_generation_state(project)
    state["regeneration_mode"] = True
    state["regeneration_from"] = target_code
    regeneration_targets = list(get_generation_regeneration_target_codes(target_code))
    state["regeneration_targets"] = regeneration_targets
    state["regeneration_job_baseline_sns"] = {
        code: getattr(find_generation_job(project, code, job_kind=GENERATION_JOB_KIND_INITIAL), "sn", 0) or 0
        for code in regeneration_targets
    }
    hydrate_generation_state_from_existing_documents(project, state)
    save_generation_state(session, state)
    return state

def get_generation_state(session, project):
    state = session.get(GENERATION_SESSION_KEY)
    if not state:
        state = _build_empty_generation_state(project)
        return hydrate_generation_state_from_existing_documents(project, state)
    if state.get("project_sn") != getattr(project, "sn", None):
        clear_generation_state(session)
        state = _build_empty_generation_state(project)
        return hydrate_generation_state_from_existing_documents(project, state)
    state.setdefault("selected_file_ids", [])
    state.setdefault("draft_documents", {})
    state.setdefault("confirmed_documents", {})
    state.setdefault("itf_reference_files", [])
    state.setdefault("regeneration_mode", False)
    state.setdefault("regeneration_targets", [])
    return hydrate_generation_state_from_existing_documents(project, state)


def save_generation_state(session, state):
    session[GENERATION_SESSION_KEY] = state
    session.modified = True


def clear_generation_state(session, project=None):
    state = session.get(GENERATION_SESSION_KEY)
    if project is None or not state or state.get("project_sn") == getattr(project, "sn", None):
        if state:
            cleanup_generation_itf_references(state)
        session.pop(GENERATION_SESSION_KEY, None)
        session.modified = True


def update_generation_selected_files(state, file_ids):
    cleanup_generation_itf_references(state)
    state["selected_file_ids"] = [str(file_id) for file_id in file_ids if str(file_id).strip()]
    state["draft_documents"] = {}
    state["confirmed_documents"] = {}
    state["itf_reference_files"] = []
    return state


def get_generation_selected_files(project, state):
    return list(
        get_project_files(
            project,
            file_ids=state.get("selected_file_ids", []),
            allowed_types=ALLOWED_GENERATION_FILE_CODES,
        )
    )


def get_project_nets(project):
    if project is None:
        return []
    return list(ProjectNet.objects.filter(project=project).order_by("sn"))


def get_generation_source_inputs(project, state, document_code):
    if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE:
        return get_generation_itf_references(state)
    if document_code == ARCHITECTURE_DOCUMENT_CODE:
        return get_project_nets(project)
    if document_code in FILE_INPUT_DOCUMENT_CODES:
        return get_generation_selected_files(project, state)
    if document_code in DERIVED_DOCUMENT_CODES:
        source_documents = []
        for prerequisite_code in get_document_prerequisite_codes(document_code):
            prerequisite_document = get_generation_saved_document(project, prerequisite_code, state)
            if prerequisite_document is not None:
                source_documents.append(prerequisite_document)
        return source_documents
    return get_generation_selected_files(project, state)


def get_generation_prerequisite_error(project, state, document_code):
    missing_prerequisite_labels = get_missing_generation_prerequisite_labels(project, state, document_code)
    if missing_prerequisite_labels:
        return f"?좏뻾 ?곗텧臾?{', '.join(missing_prerequisite_labels)}) ??λ낯???꾩슂?⑸땲??"
    if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE and not get_generation_itf_references(state):
        return "?ъ슜???명꽣?섏씠??李멸퀬 ?대?吏瑜??섎굹 ?댁긽 ?낅줈?쒗빐 二쇱꽭??"
    if document_code == ARCHITECTURE_DOCUMENT_CODE and not get_project_nets(project):
        return "?꾪궎?띿쿂 援ъ꽦?붿냼瑜??섎굹 ?댁긽 異붽???二쇱꽭??"
    if document_code in FILE_INPUT_DOCUMENT_CODES and not get_generation_selected_files(project, state):
        return "?앹꽦???ъ슜??臾몄꽌瑜?癒쇱? ?좏깮??二쇱꽭??"
    return None


def get_missing_generation_prerequisite_labels(project, state, document_code):
    return [
        get_document_label(prerequisite_code)
        for prerequisite_code in get_document_prerequisite_codes(document_code)
        if get_generation_saved_document(project, prerequisite_code, state) is None
    ]


def is_generation_dependency_ready(project, state, document_code):
    return not get_missing_generation_prerequisite_labels(project, state, document_code)


def get_current_generation_code(state, project=None, preferred_code=None):
    confirmed = state.get("confirmed_documents", {})
    draft_documents = state.get("draft_documents", {})
    if preferred_code and str(preferred_code) not in confirmed and str(preferred_code) not in draft_documents:
        if project is None or is_generation_dependency_ready(project, state, preferred_code):
            return preferred_code
    for code in get_document_code_sequence():
        if str(code) in confirmed or str(code) in draft_documents:
            continue
        if project is None or is_generation_dependency_ready(project, state, code):
            return code
    return None


def is_generation_complete(state):
    confirmed = state.get("confirmed_documents", {}) or {}
    return all(str(code) in confirmed for code in get_document_code_sequence())


def get_generation_progress_rows(state, project=None):
    state = state or {}
    draft_documents = state.get("draft_documents", {}) or {}
    confirmed_documents = state.get("confirmed_documents", {}) or {}
    rows = []

    for code in get_document_code_sequence():
        prerequisite_codes = get_document_prerequisite_codes(code)
        prerequisite_labels = [get_document_label(prerequisite_code) for prerequisite_code in prerequisite_codes]
        missing_prerequisite_labels = get_missing_generation_prerequisite_labels(project, state, code)
        dependency_ready = project is None or not missing_prerequisite_labels

        status = "pending"
        status_label = "생성 대기"
        job_status_code = PROGRESS_PENDING
        document_sn = None
        detail_url = ""
        download_url = ""

        if code in confirmed_documents or str(code) in confirmed_documents:
            status = "confirmed"
            status_label = "확정 완료"
            document_sn = confirmed_documents.get(code) or confirmed_documents.get(str(code))
            if document_sn:
                detail_url = reverse("doc_detail", args=[document_sn])
                download_url = f'{reverse("doc_content", args=[document_sn])}?download=1'
        elif code in draft_documents or str(code) in draft_documents:
            status = "review"
            status_label = "검토 중"
            document_sn = draft_documents.get(code) or draft_documents.get(str(code))
            if document_sn:
                detail_url = reverse("doc_detail", args=[document_sn])
        elif project is not None:
            latest_job = find_generation_job(project, code, job_kind=GENERATION_JOB_KIND_INITIAL)
            latest_job_status = getattr(latest_job, "job_status_id", "")
            latest_job_label = getattr(getattr(latest_job, "job_status", None), "name", "") or latest_job_status
            if latest_job_status == PROGRESS_PROCESSING:
                status = "processing"
                status_label = latest_job_label or "생성 중"
                job_status_code = latest_job_status
            elif latest_job_status == PROGRESS_FAILED:
                status = "failed"
                status_label = latest_job_label or "생성 실패"
                job_status_code = latest_job_status
            elif latest_job_status == PROGRESS_PENDING and dependency_ready:
                status = "pending"
                status_label = latest_job_label or "생성 대기"
                job_status_code = latest_job_status
            elif not dependency_ready:
                status = "locked"
                status_label = "선행 RA 문서 대기"
        elif not dependency_ready:
            status = "locked"
            status_label = "선행 RA 문서 대기"

        rows.append({
            "code": code,
            "label": get_document_label(code),
            "status": status,
            "status_label": status_label,
            "job_status_code": job_status_code,
            "document_sn": document_sn,
            "detail_url": detail_url,
            "download_url": download_url,
            "prerequisite_labels": prerequisite_labels,
            "missing_prerequisite_labels": missing_prerequisite_labels,
        })
    return rows

def build_document_detail_url(document, *, mode=None):
    query_items = []
    if mode:
        query_items.append(("mode", mode))
    base_url = reverse("doc_detail", args=[document.sn])
    return f"{base_url}?{urlencode(query_items)}" if query_items else base_url


def build_generation_steps(document_code):
    label = get_document_label(document_code)
    return [
        f"{label} ?먮룞 ?앹꽦???붿껌?덉뒿?덈떎.",
        "?좏깮???뚯씪??遺꾩꽍?섍퀬 ?붾? ?묐떟??以鍮꾪븯怨??덉뒿?덈떎.",
        "珥덉븞 援ъ“瑜??뺣━?섍퀬 ?덉뒿?덈떎.",
        f"{label} 珥덉븞??以鍮꾨릺?덉뒿?덈떎. ?뺤씤 踰꾪듉?쇰줈 寃곌낵瑜?寃?좏빐 二쇱꽭??",
    ]


@transaction.atomic
def create_project_net(
    *,
    project,
    actor,
    name,
    purpose="",
    middleware_stack="",
    firewall_settings="",
    auth_method="",
    expected_concurrent_users=None,
    cloud_yn=None,
    hardware_spec="",
    remarks="",
):
    return ProjectNet.objects.create(
        project=project,
        name=name,
        purpose=purpose or None,
        middleware_stack=middleware_stack or None,
        firewall_settings=firewall_settings or None,
        auth_method=auth_method or None,
        expected_concurrent_users=expected_concurrent_users,
        cloud_yn=cloud_yn,
        hardware_spec=hardware_spec or None,
        remarks=remarks or None,
        created_by=actor,
        updated_by=actor,
    )


def get_document_view_state(document, actor, preferred_mode="view"):
    pending_approval = get_latest_pending_approval(document)
    if pending_approval and pending_approval.created_by_id == actor.sn:
        return "waiting", pending_approval

    if document.possession_user_id and document.possession_user_id != actor.sn:
        return "readonly", pending_approval
    if document.possession_user_id and document.possession_user_id == actor.sn:
        return "edit", pending_approval
    return "view", pending_approval


def wait_for_new_revision(document, *, baseline_detail_sn=None, timeout_seconds=5, interval_seconds=0.25):
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        latest_detail = get_latest_detail(document)
        if latest_detail is not None and latest_detail.sn != baseline_detail_sn:
            return latest_detail
        time.sleep(interval_seconds)
    return get_latest_detail(document)


def build_editor_config(request, document, actor, mode, detail=None):
    latest_detail = detail or get_latest_detail(document)
    public_base_url = get_public_base_url(request)
    if not public_base_url:
        public_base_url = request.build_absolute_uri("/").rstrip("/")

    document_query_items = []
    if latest_detail is not None:
        document_query_items.append(("detail_sn", str(latest_detail.sn)))
    if settings.ONLYOFFICE_JWT_SECRET:
        document_query_items.append(
            (
                "token",
                encode_jwt(
                    {
                        "document_sn": document.sn,
                        "project_sn": document.project_id,
                        "detail_sn": getattr(latest_detail, "sn", None),
                    },
                    settings.ONLYOFFICE_JWT_SECRET,
                ),
            )
        )

    document_url = f"{public_base_url}{reverse('doc_content', args=[document.sn])}"
    if document_query_items:
        document_url = f"{document_url}?{urlencode(document_query_items)}"

    callback_url = f"{public_base_url}{reverse('doc_callback', args=[document.sn])}"
    if latest_detail is not None:
        callback_url = f"{callback_url}?{urlencode({'baseline_detail_sn': latest_detail.sn})}"

    payload = {
        "documentType": "word",
        "width": "100%",
        "document": {
            "title": get_document_title(document),
            "url": document_url,
            "fileType": "docx",
            "key": build_document_key(document, latest_detail=latest_detail),
            "permissions": {
                "edit": mode == "edit",
                "download": True,
                "print": True,
                "comment": False,
                "review": False,
            },
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": mode,
            "user": {"id": str(actor.sn), "name": actor.name},
        },
        "type": "desktop" if mode == "edit" else "embedded",
    }

    if settings.ONLYOFFICE_JWT_SECRET:
        payload["token"] = encode_jwt(payload, settings.ONLYOFFICE_JWT_SECRET)
    return payload


def parse_callback_payload(request):
    body = request.body.decode("utf-8") if request.body else "{}"
    payload = json.loads(body)

    auth_header = request.headers.get("Authorization", "")
    token = None
    if auth_header.startswith("Bearer "):
        token = auth_header.split(" ", 1)[1].strip()
    elif payload.get("token"):
        token = payload["token"]

    if settings.ONLYOFFICE_JWT_SECRET and token:
        decode_jwt(token, settings.ONLYOFFICE_JWT_SECRET)
    return payload


def validate_document_content_token(document, token, detail_sn=None):
    if not settings.ONLYOFFICE_JWT_SECRET or not token:
        return False
    try:
        payload = decode_jwt(token, settings.ONLYOFFICE_JWT_SECRET)
    except Exception:
        return False
    valid = (
        str(payload.get("document_sn")) == str(document.sn)
        and str(payload.get("project_sn")) == str(document.project_id)
    )
    if not valid:
        return False
    token_detail_sn = payload.get("detail_sn")
    if detail_sn is not None and token_detail_sn is not None:
        return str(token_detail_sn) == str(detail_sn)
    return True


def _coerce_json_value(value):
    if isinstance(value, (dict, list)):
        return value
    if not isinstance(value, str) or not value.strip():
        return value
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return value


def build_approval_data_view(value):
    data = _coerce_json_value(value)
    if isinstance(data, dict) and isinstance(data.get("tables"), list):
        tables = []
        for table in data["tables"]:
            if not isinstance(table, dict):
                continue
            columns = table.get("columns") if isinstance(table.get("columns"), list) else []
            tables.append(
                {
                    "table_id": table.get("table_id") or "-",
                    "table_name": table.get("table_name") or "",
                    "description": table.get("description") or "",
                    "database_name": table.get("database_name") or "",
                    "tablespace_name": table.get("tablespace_name") or "",
                    "columns": [column for column in columns if isinstance(column, dict)],
                }
            )
        return {"kind": "tables", "tables": tables, "empty": not tables}

    if data in (None, "", [], {}):
        return {"kind": "empty", "empty": True}

    return {
        "kind": "json",
        "formatted": json.dumps(data, ensure_ascii=False, indent=2, default=str),
        "empty": False,
    }


JSON_IDENTITY_KEYS = (
    "requirement_id",
    "component_id",
    "table_id",
    "column_id",
    "entity_id",
    "relationship_id",
    "screen_id",
    "interface_id",
    "scenario_id",
    "test_case_id",
    "id",
    "code",
    "name",
)


def _find_json_list_identity_key(before_items, after_items):
    items = [item for item in [*before_items, *after_items] if isinstance(item, dict)]
    if not items or len(items) != len(before_items) + len(after_items):
        return None
    for key in JSON_IDENTITY_KEYS:
        before_values = [str(item.get(key, "")) for item in before_items]
        after_values = [str(item.get(key, "")) for item in after_items]
        if (
            all(before_values)
            and all(after_values)
            and len(before_values) == len(set(before_values))
            and len(after_values) == len(set(after_values))
        ):
            return key
    return None


def _build_generic_json_changes(before, after, *, path="data", changes=None, limit=200):
    changes = [] if changes is None else changes
    if len(changes) >= limit:
        return changes

    if type(before) is not type(after):
        changes.append(_build_generic_change(path, "modified", before, after))
        return changes

    if isinstance(before, dict):
        for key in sorted(set(before) | set(after)):
            child_path = f"{path}.{key}"
            if key not in before:
                changes.append(_build_generic_change(child_path, "added", None, after[key]))
            elif key not in after:
                changes.append(_build_generic_change(child_path, "deleted", before[key], None))
            else:
                _build_generic_json_changes(before[key], after[key], path=child_path, changes=changes, limit=limit)
            if len(changes) >= limit:
                break
        return changes

    if isinstance(before, list):
        identity_key = _find_json_list_identity_key(before, after)
        if identity_key:
            before_map = {str(item[identity_key]): item for item in before}
            after_map = {str(item[identity_key]): item for item in after}
            for identity in sorted(set(before_map) | set(after_map)):
                child_path = f"{path}[{identity_key}={identity}]"
                if identity not in before_map:
                    changes.append(_build_generic_change(child_path, "added", None, after_map[identity]))
                elif identity not in after_map:
                    changes.append(_build_generic_change(child_path, "deleted", before_map[identity], None))
                else:
                    _build_generic_json_changes(
                        before_map[identity],
                        after_map[identity],
                        path=child_path,
                        changes=changes,
                        limit=limit,
                    )
                if len(changes) >= limit:
                    break
            return changes

        for index in range(max(len(before), len(after))):
            child_path = f"{path}[{index}]"
            if index >= len(before):
                changes.append(_build_generic_change(child_path, "added", None, after[index]))
            elif index >= len(after):
                changes.append(_build_generic_change(child_path, "deleted", before[index], None))
            else:
                _build_generic_json_changes(before[index], after[index], path=child_path, changes=changes, limit=limit)
            if len(changes) >= limit:
                break
        return changes

    if before != after:
        changes.append(_build_generic_change(path, "modified", before, after))
    return changes


def _build_generic_change(path, change_type, before, after):
    title = path.rsplit(".", 1)[-1]
    message_by_type = {
        "added": f"{title} ??ぉ??異붽??섏뿀?듬땲??",
        "deleted": f"{title} ??ぉ????젣?섏뿀?듬땲??",
        "modified": f"{title} 媛믪씠 蹂寃쎈릺?덉뒿?덈떎.",
    }
    return {
        "title": title,
        "target_path": path,
        "change_type": change_type,
        "before": before,
        "after": after,
        "message": message_by_type[change_type],
        "classification_source": "json_diff_fallback",
    }


def build_approval_review_view(value, *, before_data=None, after_data=None):
    data = _coerce_json_value(value)
    if not isinstance(data, dict):
        data = {}

    change_review = data.get("change_review") if isinstance(data.get("change_review"), dict) else {}
    consistency = data.get("consistency_check") if isinstance(data.get("consistency_check"), dict) else {}
    changes = change_review.get("changes") if isinstance(change_review.get("changes"), list) else []
    changes = [change for change in changes if isinstance(change, dict)]
    if not changes and before_data is not None and after_data is not None:
        changes = _build_generic_json_changes(
            _coerce_json_value(before_data),
            _coerce_json_value(after_data),
        )
    messages = consistency.get("messages") if isinstance(consistency.get("messages"), list) else []
    normalized_changes = []
    for change in changes:
        normalized_change = dict(change)
        normalized_change["before_display"] = _format_approval_change_value(change.get("before"))
        normalized_change["after_display"] = _format_approval_change_value(change.get("after"))
        normalized_changes.append(normalized_change)
    change_summary = change_review.get("summary") if isinstance(change_review.get("summary"), dict) else {}
    if changes and not any(change_summary.get(key) for key in ("added_count", "deleted_count", "modified_count")):
        change_summary = {
            "added_count": sum(change.get("change_type") == "added" for change in changes),
            "deleted_count": sum(change.get("change_type") == "deleted" for change in changes),
            "modified_count": sum(change.get("change_type") == "modified" for change in changes),
        }
    return {
        "status": data.get("status") or consistency.get("status") or "",
        "change_summary": change_summary,
        "changes": normalized_changes,
        "consistency_summary": consistency.get("summary") if isinstance(consistency.get("summary"), dict) else {},
        "consistency_messages": [message for message in messages if isinstance(message, dict)],
    }


def _format_approval_change_value(value):
    if value is None:
        return "없음"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2, default=str)
    if isinstance(value, bool):
        return "예" if value else "아니오"
    return str(value)


def get_project_files(project, *, file_ids=None, allowed_types=None):
    queryset = ProjectFile.objects.filter(project=project).select_related("file_type", "created_by")
    if allowed_types:
        queryset = queryset.filter(file_type_id__in=allowed_types)
    if file_ids:
        queryset = queryset.filter(sn__in=file_ids)
    return queryset.order_by("-created_at", "-sn")


def build_document_rows(queryset):
    rows = []
    for document in queryset:
        detail_url = f'{reverse("doc_detail", args=[document.sn])}?from=history'
        rows.append(
            {
                "sn": document.sn,
                "type_name": getattr(document.document_type, "name", "-") or "-",
                "creator_name": getattr(document.created_by, "name", "-") or "-",
                "version": document.version or "-",
                "modification_content": document.modification_content or "-",
                "created_at": document.created_at,
                "detail_url": detail_url,
                "download_url": f'{reverse("doc_content", args=[document.sn])}?download=1',
                "locked_by_name": getattr(document.possession_user, "name", ""),
                "status_label": "문서 보관",
            }
        )
    return rows


def build_approval_rows(queryset):
    rows = []
    for approval in queryset:
        rows.append(
            {
                "sn": approval.approval_sn,
                "document_sn": approval.detail.document.sn,
                "document_label": approval.detail.document.document_type.name,
                "version": approval.detail.document.version,
                "requester_name": getattr(approval.created_by, "name", "-") or "-",
                "status_name": approval.approval_status.name,
                "status_code": approval.approval_status_id,
                "request_content": approval.request_content,
                "created_at": approval.created_at,
                "detail_url": reverse("doc_approval_detail", args=[approval.approval_sn]),
            }
        )
    return rows


def build_approval_queryset(project, actor):
    queryset = (
        DocumentApproval.objects.filter(detail__document__project=project)
        .select_related(
            "detail__document__document_type",
            "approval_status",
            "created_by",
            "detail__document__project",
        )
        .order_by("-created_at", "-approval_sn")
    )
    if not is_project_manager(project, actor):
        queryset = queryset.filter(created_by=actor)
    return queryset


def apply_approval_filters(params, queryset, *, include_requester=True):
    document_code = params.get("docs_cd", "all")
    approval_status = params.get("status", "all")
    requester_query = params.get("requester", "").strip()

    if document_code != "all":
        queryset = queryset.filter(detail__document__document_type_id=document_code)
    if approval_status != "all":
        queryset = queryset.filter(approval_status_id=approval_status)
    if include_requester and requester_query:
        queryset = queryset.filter(created_by__name__icontains=requester_query)

    return queryset, document_code, approval_status, requester_query



