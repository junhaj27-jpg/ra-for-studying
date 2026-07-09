"""RFP 문서에서 사용자 요구사항 정의서 입력 JSON을 추출하는 Rule Parser입니다."""

import copy
import json
import re
import zipfile
from collections.abc import Callable
from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from tools.result import ToolResult, error_result, success_result
from tools.parser.pdf_parser import parse_pdf


RuleParser = Callable[[str], Any]

ID_PATTERN = re.compile(r"^[A-Za-z]{2,10}[\-\u2013\u2014_]\d{1,5}$")

DEFAULT_PREFIX_MAP = {
    "SFR": "기능",
    "FUR": "기능",
    "ECR": "시스템 장비구성",
    "NFR": "비기능",
    "PER": "성능",
    "SER": "보안",
    "SEC": "보안",
    "QUR": "품질",
    "TER": "테스트",
    "TST": "테스트",
    "DAR": "데이터",
    "DBR": "데이터",
    "CSR": "컨설팅",
    "PSR": "프로젝트 지원",
    "PMR": "프로젝트 관리",
    "COR": "제약사항",
    "INT": "인터페이스",
    "INR": "인터페이스",
    "SIR": "시스템 장비구성",
    "MPR": "유지보수",
}

FIELD_ALIASES = {
    "name": {
        "요구사항명",
        "요구사항 명",
        "요구사항 명칭",
        "기능명",
        "기능 명",
        "항목명",
        "업무명",
    },
    "type": {"요구사항분류", "요구사항 구분", "구분", "분류", "유형"},
    "definition": {"요구사항 정의", "요구사항정의", "정의"},
    "detail": {"요구사항 상세설명", "상세설명", "상세 설명", "세부내용", "세부 내용"},
    "constraint": {"제약사항", "특이사항", "조건", "비고"},
    "validation": {"검증기준", "시험기준", "평가기준", "검수기준", "검토기준"},
    "priority": {"우선순위", "중요도"},
    "source_ref": {"관련 요구사항", "관련요구사항", "출처", "근거"},
}

BLACKLIST = {
    "요구사항명",
    "요구사항 명",
    "요구사항 명칭",
    "요구사항 고유번호",
    "요구사항 상세설명",
    "상세설명",
    "상세 설명",
    "세부내용",
    "세부 내용",
    "정의",
    "산출정보",
    "관련 요구사항",
    "관련요구사항",
    "검토기준",
    "검증기준",
    "비고",
    "구분",
    "분류",
    "유형",
    "사업수행계획서",
    "설계서",
    "결과보고서",
}


def parse_rfp_requirements(
    file_path: str,
    *,
    parser: RuleParser | None = None,
    document_id: str = "DOC-001",
) -> ToolResult:
    """RFP 파일에서 문서 메타데이터와 기능 요구사항 목록을 추출합니다."""

    selected_parser = parser or extract_requirements_from_rfp
    try:
        requirements = selected_parser(file_path)
        if isinstance(requirements, dict):
            requirements = (
                requirements.get("requirements")
                or requirements.get("requirement_json_list")
                or requirements.get("functional_requirements")
                or []
            )
        if not isinstance(requirements, list):
            raise ValueError("파서 결과는 요구사항 목록이어야 합니다.")
        if not requirements:
            raise ValueError("RFP 문서에서 요구사항 항목을 추출하지 못했습니다.")
        functional_requirements = [
            requirement
            for requirement in requirements
            if isinstance(requirement, dict)
            and _is_functional_type(requirement.get("requirement_type"))
        ]
        non_functional_requirements = [
            requirement
            for requirement in requirements
            if isinstance(requirement, dict)
            and not _is_functional_type(requirement.get("requirement_type"))
        ]
        return success_result(
            {
                "document_id": document_id,
                "document_name": Path(file_path).name,
                "requirements": [requirement for requirement in requirements if isinstance(requirement, dict)],
                "functional_requirements": functional_requirements,
                "non_functional_requirements": non_functional_requirements,
            }
        )
    except Exception as exc:
        return error_result("RFP_RULE_PARSE_FAILED", str(exc), {"file_path": file_path})


def extract_requirements_from_rfp(file_path: str) -> list[dict[str, Any]]:
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        return extract_requirements_from_rfp_pdf(file_path)
    return extract_requirements_from_rfp_docx(file_path)


def _is_functional_type(value: Any) -> bool:
    requirement_type = str(value or "").strip().lower()
    return requirement_type.startswith("기능") or requirement_type.startswith("functional") or requirement_type == "function"


def extract_requirements_from_rfp_docx(file_path: str) -> list[dict[str, Any]]:
    """DOCX 표를 순회하며 요구사항 ID 기준으로 항목을 구성합니다."""

    path = Path(file_path)
    with _normalized_docx_package(path) as readable_path:
        document = Document(readable_path)
    requirements_map: dict[str, dict[str, Any]] = {}

    for table_index, table in enumerate(document.tables):
        last_id: str | None = None
        for row_index, row in enumerate(table.rows):
            cells = _unique_cells(row.cells)
            if not cells:
                continue

            requirement_id = _find_requirement_id(cells)
            if requirement_id:
                last_id = requirement_id
                current = requirements_map.setdefault(
                    requirement_id,
                    _new_requirement(requirement_id, path.name, table_index, row_index),
                )
                _apply_id_row_values(current, cells, requirement_id)
                continue

            if not last_id:
                continue

            current = requirements_map[last_id]
            if len(cells) >= 2:
                field_value = _split_docx_field(cells)
                field, value = field_value if field_value else (None, "")
                if field == "name":
                    current["name"] = value
                    continue
                if field == "type":
                    current["type"] = value
                    continue
                if field == "definition":
                    current["definition_parts"].append(value)
                    _update_source_location(current, table_index, row_index)
                    continue
                if field == "detail":
                    current["detail_parts"].append(value)
                    _update_source_location(current, table_index, row_index)
                    continue
                if field == "constraint":
                    current["constraints"].append(value)
                    continue
                if field == "validation":
                    current["validation_criteria"].append(value)
                    continue
                if field == "priority":
                    current["priority"] = value
                    continue
                if field == "source_ref":
                    current["source_refs"].append(value)
                    continue
                if "산출정보" in cells[0]:
                    continue

            current["detail_parts"].extend(cell for cell in cells if cell not in BLACKLIST)

    table_requirements = [
        _build_requirement(data)
        for data in requirements_map.values()
        if _is_valid_requirement(data)
    ]
    if table_requirements:
        return table_requirements

    text_lines = _docx_text_lines(document)
    return _extract_requirements_from_lines(
        text_lines,
        source_name=path.name,
        source_type="detailed_requirement_text",
    )


@contextmanager
def _normalized_docx_package(path: Path):
    """Linux에서 열 수 없는 역슬래시 ZIP 경로를 메모리에서 표준화합니다.

    DOCX는 ZIP 내부 경로에 POSIX 구분자(`/`)를 사용해야 합니다. 일부
    Windows 생성기가 `word\\document.xml`, `_rels\\.rels` 형태로 저장한
    파일은 Windows에서는 열리지만 Linux의 python-docx에서는 실패합니다.
    이 경우에만 원본 내용은 그대로 두고 엔트리 경로만 바꾼 메모리 버퍼를 사용합니다.
    """

    normalized_buffer: BytesIO | None = None
    try:
        with zipfile.ZipFile(path) as source:
            members = source.infolist()
            if not any("\\" in member.filename for member in members):
                yield path
                return

            normalized_buffer = BytesIO()
            normalized_names: set[str] = set()
            with zipfile.ZipFile(normalized_buffer, "w") as target:
                for member in members:
                    normalized_name = member.filename.replace("\\", "/")
                    if normalized_name in normalized_names:
                        raise ValueError(
                            f"DOCX 경로 표준화 중 중복 엔트리가 발생했습니다: {normalized_name}"
                        )
                    normalized_names.add(normalized_name)
                    normalized_member = copy.copy(member)
                    normalized_member.filename = normalized_name
                    target.writestr(normalized_member, source.read(member.filename))

        normalized_buffer.seek(0)
        yield normalized_buffer
    finally:
        if normalized_buffer is not None:
            normalized_buffer.close()


def extract_requirements_from_rfp_pdf(file_path: str) -> list[dict[str, Any]]:
    """PDF 텍스트에서 요구사항 ID 블록을 찾아 요구사항 항목을 구성합니다."""

    parsed = parse_pdf(file_path)
    if not parsed["success"]:
        error = parsed["error"] or {}
        raise ValueError(str(error.get("message", "PDF 파싱에 실패했습니다.")))

    path = Path(file_path)
    requirements_map: dict[str, dict[str, Any]] = {}
    for page in parsed["data"].get("pages", []):
        page_number = int(page.get("page_number") or 0)
        lines = _pdf_lines(str(page.get("text") or ""))
        current_id: str | None = None
        line_index = 0
        while line_index < len(lines):
            line = lines[line_index]
            requirement_id = _find_requirement_id_in_text(line)
            if requirement_id:
                current_id = requirement_id
                current = requirements_map.setdefault(
                    requirement_id,
                    _new_requirement(
                        requirement_id,
                        path.name,
                        page_number,
                        line_index,
                        source_type="detailed_requirement_text",
                    ),
                )
                rest = _remove_requirement_id_label(line, requirement_id)
                if rest:
                    current["detail_parts"].append(rest)
                line_index += 1
                continue

            if not current_id:
                line_index += 1
                continue

            current = requirements_map[current_id]
            field_value = _split_pdf_field(line)
            if field_value:
                field, value = field_value
                _apply_pdf_field(current, field, value)
                line_index += 1
                continue

            field = _detect_field(line)
            if field:
                value, next_index = _collect_pdf_field_value(lines, line_index + 1, field)
                if value:
                    _apply_pdf_field(current, field, value)
                    line_index = next_index
                else:
                    line_index += 1
                continue

            if _looks_like_section_boundary(line):
                line_index += 1
                continue
            current["detail_parts"].append(line)
            line_index += 1

    return [_build_requirement(data) for data in requirements_map.values() if _is_valid_requirement(data)]


def _docx_text_lines(document: Any) -> list[str]:
    lines: list[str] = []
    lines.extend(
        normalize_text(paragraph.text)
        for paragraph in document.paragraphs
        if normalize_text(paragraph.text)
    )
    for table in document.tables:
        for row in table.rows:
            cells = _unique_cells(row.cells)
            if cells:
                lines.append(" ".join(cells))
    return [line for line in lines if line]


def _extract_requirements_from_lines(
    lines: list[str],
    *,
    source_name: str,
    source_type: str,
) -> list[dict[str, Any]]:
    requirements_map: dict[str, dict[str, Any]] = {}
    current_id: str | None = None
    line_index = 0
    while line_index < len(lines):
        line = lines[line_index]
        requirement_id = _find_requirement_id_in_text(line)
        if requirement_id:
            current_id = requirement_id
            current = requirements_map.setdefault(
                requirement_id,
                _new_requirement(
                    requirement_id,
                    source_name,
                    0,
                    line_index,
                    source_type=source_type,
                ),
            )
            rest = _remove_requirement_id_label(line, requirement_id)
            if rest:
                field_value = _split_pdf_field(rest)
                if field_value:
                    _apply_pdf_field(current, field_value[0], field_value[1])
                else:
                    current["detail_parts"].append(rest)
            line_index += 1
            continue

        if not current_id:
            line_index += 1
            continue

        current = requirements_map[current_id]
        field_value = _split_pdf_field(line)
        if field_value:
            field, value = field_value
            _apply_pdf_field(current, field, value)
            line_index += 1
            continue

        field = _detect_field(line)
        if field:
            value, next_index = _collect_pdf_field_value(lines, line_index + 1, field)
            if value:
                _apply_pdf_field(current, field, value)
                line_index = next_index
            else:
                line_index += 1
            continue

        if _looks_like_section_boundary(line):
            line_index += 1
            continue
        current["detail_parts"].append(line)
        line_index += 1

    return [
        _build_requirement(data)
        for data in requirements_map.values()
        if _is_valid_requirement(data)
    ]


def _new_requirement(
    requirement_id: str,
    file_name: str,
    table_index: int,
    row_index: int,
    *,
    source_type: str = "detailed_requirement_table",
) -> dict[str, Any]:
    return {
        "id": requirement_id,
        "name": None,
        "type": None,
        "priority": None,
        "constraints": [],
        "validation_criteria": [],
        "source_refs": [],
        "definition_parts": [],
        "detail_parts": [],
        "source": file_name,
        "table_index": table_index,
        "row_index": row_index,
        "source_type": source_type,
    }


def _build_requirement(data: dict[str, Any]) -> dict[str, Any]:
    req_id = data["id"]
    definition_parts = _clean_parts(data["definition_parts"])
    detail_parts = _clean_parts(data["detail_parts"])
    name = data["name"] or _infer_requirement_name([*definition_parts, *detail_parts])
    definition = _clean_field_text("\n".join(definition_parts))
    detail = _clean_field_text("\n".join(detail_parts))
    req_type = _normalize_requirement_type(
        req_id,
        data.get("type"),
        name,
        detail,
    )

    return {
        "requirement_id": req_id,
        "requirement_name": name[:100],
        "requirement_type": req_type,
        "requirement_definition": definition,
        "requirement_detail": detail,
        "source_location": {
            "table_index": data["table_index"],
            "source_type": data["source_type"],
        },
    }


def _unique_cells(cells: Any) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for cell in cells:
        text = normalize_multiline_text(cell.text)
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def _apply_id_row_values(
    current: dict[str, Any],
    cells: list[str],
    requirement_id: str,
) -> None:
    """요구사항 목록형 표의 ID 행에서 유형과 명칭만 보강합니다."""

    values = [
        cell
        for cell in cells
        if _normalize_id(cell) != requirement_id
        and cell not in BLACKLIST
        and cell not in {"ID", "요구사항 번호"}
    ]
    if len(values) >= 2:
        current["type"] = current.get("type") or values[0]
        current["name"] = current.get("name") or values[-1]


def _split_docx_field(cells: list[str]) -> tuple[str, str] | None:
    """병합 셀이 포함된 DOCX 행에서 실제 필드와 값을 분리합니다."""

    for index in range(len(cells) - 2, -1, -1):
        cell = cells[index]
        field = _detect_field(cell)
        if field in {"definition", "detail"}:
            value = " ".join(cells[index + 1 :]).strip()
            if value:
                return field, value

    field = _detect_field(cells[0])
    value = " ".join(cells[1:]).strip()
    return (field, value) if field and value else None


def _update_source_location(
    current: dict[str, Any],
    table_index: int,
    row_index: int,
) -> None:
    current["table_index"] = table_index
    current["row_index"] = row_index
    current["source_type"] = "detailed_requirement_table"


def _find_requirement_id(cells: list[str]) -> str | None:
    for cell in cells:
        normalized = _normalize_id(cell)
        if ID_PATTERN.match(normalized):
            return normalized
    return None


def _find_requirement_id_in_text(text: str) -> str | None:
    normalized = _normalize_id(text)
    match = re.search(r"\b[A-Z]{2,10}-\d{1,5}\b", normalized)
    return match.group(0) if match else None


def _normalize_id(text: str) -> str:
    value = re.sub(r"[\-\u2013\u2014_]", "-", text.strip()).upper()
    value = re.sub(r"\b([A-Z]{2,10})\s+(\d{1,5})\b", r"\1-\2", value)
    return value


def normalize_text(text: str) -> str:
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_multiline_text(text: str) -> str:
    lines = [normalize_text(line) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _detect_field(text: str) -> str | None:
    normalized = normalize_text(text)
    for field, aliases in FIELD_ALIASES.items():
        if any(normalized == alias or alias in normalized for alias in aliases):
            return field
    return None


def _pdf_lines(text: str) -> list[str]:
    return [normalize_text(line) for line in text.splitlines() if normalize_text(line)]


def _remove_requirement_id_label(text: str, requirement_id: str) -> str:
    value = _normalize_id(text)
    value = value.replace(requirement_id, " ")
    value = re.sub(r"(요구사항\s*고유번호|요구사항\s*번호|요구사항\s*ID|ID)\s*[:：]?", " ", value, flags=re.IGNORECASE)
    return normalize_text(value)


def _split_pdf_field(text: str) -> tuple[str, str] | None:
    normalized = normalize_text(text)
    if any(normalized == alias for aliases in FIELD_ALIASES.values() for alias in aliases):
        return None
    for separator in (":", "："):
        if separator in normalized:
            label, value = [part.strip() for part in normalized.split(separator, 1)]
            field = _detect_field(label)
            if field and value:
                return field, value

    field = _detect_field(normalized)
    if not field:
        return None
    for aliases in FIELD_ALIASES.values():
        for alias in sorted(aliases, key=len, reverse=True):
            if normalized.startswith(alias):
                value = normalize_text(normalized[len(alias):])
                value = re.sub(r"^[\s\-:：]+", "", value)
                if value:
                    return field, value
    return None


def _collect_pdf_field_value(
    lines: list[str],
    start_index: int,
    field: str,
) -> tuple[str, int]:
    values: list[str] = []
    index = start_index
    while index < len(lines):
        line = lines[index]
        if _find_requirement_id_in_text(line):
            break
        if _detect_field(line):
            break
        if line in BLACKLIST or line in {"요구사항", "상세설명"}:
            index += 1
            if field in {"name", "type", "priority", "source_ref"} and values:
                break
            continue
        if _looks_like_section_boundary(line):
            if values:
                break
            index += 1
            continue
        values.append(line)
        index += 1
        if field in {"type", "priority", "source_ref", "validation", "constraint"}:
            break
        if field == "name" and len(values) >= 4:
            break
        if field == "detail" and len(values) >= 20:
            break
    return _clean_pdf_field_value(field, values), index


def _clean_pdf_field_value(field: str, values: list[str]) -> str:
    if not values:
        return ""
    if field == "name":
        value = " ".join(values)
        value = re.sub(r"\s*([()])\s*", r"\1", value)
        value = re.sub(r"\s+", " ", value)
        return value.strip()
    return "\n".join(values).strip()


def _apply_pdf_field(current: dict[str, Any], field: str, value: str) -> None:
    if not value:
        return
    if field == "name":
        current["name"] = value
    elif field == "type":
        current["type"] = value
    elif field == "definition":
        current["definition_parts"].append(value)
    elif field == "detail":
        current["detail_parts"].append(value)
    elif field == "constraint":
        current["constraints"].append(value)
    elif field == "validation":
        current["validation_criteria"].append(value)
    elif field == "priority":
        current["priority"] = value
    elif field == "source_ref":
        current["source_refs"].append(value)


def _looks_like_section_boundary(text: str) -> bool:
    normalized = normalize_text(text)
    if normalized in BLACKLIST:
        return True
    return bool(re.match(r"^\d+(\.\d+){0,4}\s+.{1,30}$", normalized))


def _clean_parts(parts: list[str]) -> list[str]:
    unique_parts: list[str] = []
    seen: set[str] = set()
    for part in parts:
        text = normalize_multiline_text(part)
        if not text or text in BLACKLIST or len(text) <= 2 or text in seen:
            continue
        seen.add(text)
        unique_parts.append(text)
    return unique_parts


def _clean_field_text(text: str) -> str:
    cleaned_lines = []
    for line in text.splitlines():
        cleaned = re.sub(
            r"^(요구사항\s*고유번호|요구사항\s*상세설명|세부\s*내용|정의)\s*[:：]?\s*",
            "",
            line,
        ).strip()
        if cleaned:
            cleaned_lines.append(cleaned)
    return "\n".join(cleaned_lines).strip()


def _is_valid_requirement(data: dict[str, Any]) -> bool:
    req_id = data["id"]
    if req_id.endswith("-000") or req_id.endswith("-00"):
        return False
    content = "\n".join(
        [
            *_clean_parts(data["definition_parts"]),
            *_clean_parts(data["detail_parts"]),
        ]
    )
    return len(_clean_field_text(content)) > 20


def _infer_requirement_name(parts: list[str]) -> str:
    for part in parts:
        text = normalize_text(part)
        if len(text) < 4 or text.isdigit() or text in BLACKLIST:
            continue
        if "요구사항" in text and len(text) < 20:
            continue
        return text[:100]
    return "미분류"


def _infer_requirement_type(req_id: str, name: str, description: str) -> str:
    name_text = normalize_text(name)
    if "보안" in name_text:
        return "보안"
    if "성능" in name_text:
        return "성능"
    if "품질" in name_text:
        return "품질"
    if "인터페이스" in name_text or "UI" in name_text.upper():
        return "인터페이스"
    if "데이터" in name_text:
        return "데이터"
    if "프로젝트 지원" in name_text:
        return "프로젝트 지원"
    if "지원" in name_text:
        return "프로젝트 지원"
    if "장비구성" in name_text or "인프라" in name_text:
        return "시스템 장비구성"

    prefix = req_id.split("-")[0].upper()
    if prefix in DEFAULT_PREFIX_MAP:
        return DEFAULT_PREFIX_MAP[prefix]

    text = f"{name} {description[:500]}".lower()
    if any(keyword in text for keyword in ("보안", "암호화", "접근제어", "접근 통제", "개인정보")):
        return "보안"
    if any(keyword in text for keyword in ("처리량", "응답속도", "동시접속", "throughput")):
        return "성능"
    if any(keyword in text for keyword in ("인프라", "서버", "아키텍처")):
        return "시스템 장비구성"
    return "기능"


def _normalize_requirement_type(
    req_id: str,
    raw_type: Any,
    name: str,
    description: str,
) -> str:
    prefix = req_id.split("-")[0].upper()
    inferred = _infer_requirement_type(req_id, name, description)
    if prefix in DEFAULT_PREFIX_MAP:
        default = DEFAULT_PREFIX_MAP[prefix]
    else:
        default = inferred
    if prefix in {"SFR", "FUR"}:
        return "기능"

    value = normalize_text(str(raw_type or ""))
    if not value:
        return default
    if value in {"기능", "기능 요구사항"}:
        return "기능"
    if value in {"비기능", "비기능 요구사항"}:
        return "비기능"
    if any(keyword in value for keyword in ("보안", "성능", "품질", "데이터", "인터페이스")):
        return _infer_requirement_type(req_id, value, description)
    if "시스템" in value or "장비" in value or "인프라" in value:
        return "시스템 장비구성"
    if len(value) > 20 or value.startswith(("ㅇ", "-", ",")):
        return default
    return default if prefix in DEFAULT_PREFIX_MAP else value


def dump_requirements_json(
    requirements: list[dict[str, Any]],
    *,
    document_id: str = "DOC-001",
    document_name: str = "",
) -> str:
    """수동 검증과 CLI 출력에 쓰는 JSON 직렬화 helper입니다."""

    return json.dumps(
        {
            "document_id": document_id,
            "document_name": document_name,
            "functional_requirements": requirements,
        },
        ensure_ascii=False,
        indent=2,
    )
