from pathlib import Path

from tools.result import ToolResult, error_result, success_result


def parse_docx(file_path: str) -> ToolResult:
    try:
        from docx import Document

        path = Path(file_path).resolve(strict=True)
        document = Document(str(path))
        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        tables = [
            [[cell.text.strip() for cell in row.cells] for row in table.rows]
            for table in document.tables
        ]
        return success_result(
            {
                "file_path": str(path),
                "text": "\n".join(paragraphs),
                "paragraphs": paragraphs,
                "tables": tables,
            }
        )
    except Exception as exc:
        return error_result("DOCX_PARSE_FAILED", str(exc), {"file_path": file_path})
